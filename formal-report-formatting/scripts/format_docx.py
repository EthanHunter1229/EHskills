#!/usr/bin/env python3
"""
format_docx.py

Reformats a Chinese IPO-prospectus-style Word document according to a JSON
rules file (see assets/default_rules.json). Applies page setup, heading-level
fonts/sizes/alignment (detected via regex patterns), body paragraph
justification/indent/line-spacing, three-line table borders + alignment,
and a page-number footer.

Usage:
    python3 format_docx.py <input.docx> <output.docx> --rules <rules.json>
"""
import argparse
import json
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HEADING_ORDER = ["h1", "h2", "h3", "h4"]


def load_rules(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def set_run_font(run, east_asian_name, western_name, size_pt=None, bold=None, color_hex=None):
    """Set separate East-Asian / Western typefaces on one run (Word renders
    CJK glyphs with the eastAsia font and ASCII glyphs with the ascii/hAnsi
    font automatically -- no need to split the run by script).

    color_hex is applied explicitly (default black in the rules file) because
    many Word templates bake a theme color (e.g. blue) into their built-in
    Heading styles; without overriding it, reformatted headings keep that
    color even after we've changed the font/size, which looks inconsistent
    next to the plain black body text expected in formal documents."""
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if western_name:
        run.font.name = western_name
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    if western_name:
        rFonts.set(qn("w:ascii"), western_name)
        rFonts.set(qn("w:hAnsi"), western_name)
    if east_asian_name:
        rFonts.set(qn("w:eastAsia"), east_asian_name)


STYLE_HEADING_MAP = [
    ("Title", "cover"),
    ("Heading 1", "h1"),
    ("Heading 2", "h2"),
    ("Heading 3", "h3"),
    ("Heading 4", "h4"),
]


def classify_heading(text, rules, para_index, first_nonempty_index, style_name=None):
    """Classify a paragraph into cover/h1-h4/body (None), and report whether
    the paragraph's own text already carries a number (True) or was only
    identified via its Word style name (False, meaning it has no visible
    number/label of its own -- see needs_auto_number()).

    Numbered-heading conventions (第X节 / 一、 / （一） / 1.) used by IPO
    prospectus documents are checked first since they are the most specific
    signal. If nothing matches -- e.g. the document instead relies on Word's
    built-in Heading 1/2/3 styles, as most ordinary business reports do --
    fall back to the paragraph's style name so the skill still works on
    documents that were not written with the 招股说明书 numbering convention.

    Returns a (level, has_own_numbering) tuple, or (None, False) for body text.
    """
    text = text.strip()
    if not text:
        return None, False
    # Cover title: either the very first non-empty paragraph in the document,
    # or a short standalone line mentioning "招股说明书" near the top. The
    # length cap keeps this from misfiring on ordinary body paragraphs that
    # merely mention the term in passing (e.g. "本招股说明书中，除非另有说明...").
    if para_index == first_nonempty_index:
        return "cover", True
    if "招股说明书" in text and len(text) < 80 and para_index <= first_nonempty_index + 3:
        return "cover", True
    for level in HEADING_ORDER:
        pattern = rules["heading_levels"][level]["pattern"]
        if re.match(pattern, text):
            return level, True
    if style_name:
        for prefix, level in STYLE_HEADING_MAP:
            if style_name.startswith(prefix):
                return level, False
    return None, False


CHINESE_DIGITS = "零一二三四五六七八九"


def to_chinese_numeral(n):
    """Convert a small positive int to a Chinese numeral (1-99), matching the
    style used by 一、二、三... and （一）（二）（三） heading conventions."""
    if n <= 10:
        return "十" if n == 10 else CHINESE_DIGITS[n]
    if n < 20:
        return "十" + CHINESE_DIGITS[n - 10]
    tens, ones = divmod(n, 10)
    return CHINESE_DIGITS[tens] + "十" + (CHINESE_DIGITS[ones] if ones else "")


def make_heading_number(level, n):
    """Build an auto-generated number/label in the same visual convention as
    this level's own pattern, so headings picked up only via Word style
    (no number of their own) end up looking consistent with ones that do."""
    if level == "h1":
        return f"第{to_chinese_numeral(n)}节 "
    if level == "h2":
        return f"{to_chinese_numeral(n)}、"
    if level == "h3":
        return f"（{to_chinese_numeral(n)}）"
    if level == "h4":
        return f"{n}. "
    return ""


def is_list_paragraph(paragraph):
    """True for bullet/numbered list paragraphs, which should keep their own
    list indentation and left alignment rather than being forced into the
    justified + first-line-indent body style meant for prose paragraphs."""
    style_name = paragraph.style.name if paragraph.style else None
    if style_name and "List" in style_name:
        return True
    pPr = paragraph._p.pPr
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    return False


def _cover_cfg(rules):
    cover = rules.get("heading_levels", {}).get("cover", {})
    return {
        "size_pt": cover.get("size_pt", 22),
        "font_east_asian": cover.get("font_east_asian", "黑体"),
        "font_western": cover.get("font_western", "Times New Roman"),
        "bold": True,
        "alignment": "center",
        "color_hex": cover.get("color_hex", "000000"),
    }


def apply_heading_style(paragraph, cfg):
    paragraph.alignment = ALIGN_MAP.get(cfg.get("alignment", "left"))
    pf = paragraph.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = None
    for run in paragraph.runs:
        set_run_font(
            run,
            cfg["font_east_asian"],
            cfg["font_western"],
            cfg["size_pt"],
            cfg.get("bold", False),
            cfg.get("color_hex"),
        )


def apply_body_style(paragraph, cfg, is_list=False):
    if is_list:
        # Leave the list's own numbering/indentation alone -- forcing a
        # first-line indent and justification on top of bullet indent looks
        # broken (double-indented, ragged bullets). Still apply the body
        # font/size/line-spacing so bullets match the rest of the prose.
        paragraph.alignment = ALIGN_MAP.get("left")
        pf = paragraph.paragraph_format
        pf.first_line_indent = None
    else:
        paragraph.alignment = ALIGN_MAP.get(cfg.get("alignment", "justify"))
        pf = paragraph.paragraph_format
        indent_chars = cfg.get("first_line_indent_chars", 0)
        if indent_chars:
            pf.first_line_indent = Pt(cfg["size_pt"] * indent_chars)
        else:
            pf.first_line_indent = None
    line_spacing_pt = cfg.get("line_spacing_pt")
    if line_spacing_pt:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(cfg.get("space_before_pt", 0))
    pf.space_after = Pt(cfg.get("space_after_pt", 0))
    for run in paragraph.runs:
        set_run_font(
            run,
            cfg["font_east_asian"],
            cfg["font_western"],
            cfg["size_pt"],
            cfg.get("bold", False),
            cfg.get("color_hex"),
        )


def set_page_setup(doc, rules):
    cfg = rules["page"]
    for section in doc.sections:
        section.page_width = Cm(cfg["page_width_cm"])
        section.page_height = Cm(cfg["page_height_cm"])
        section.top_margin = Cm(cfg["margin_top_cm"])
        section.bottom_margin = Cm(cfg["margin_bottom_cm"])
        section.left_margin = Cm(cfg["margin_left_cm"])
        section.right_margin = Cm(cfg["margin_right_cm"])


def _set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)
    return borders


def apply_three_line_border(table):
    borders = _set_table_borders_none(table)

    def set_edge(name, sz):
        el = borders.find(qn(f"w:{name}"))
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    set_edge("top", 12)
    set_edge("bottom", 12)

    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
            tcBorders.append(bottom)
            tcPr.append(tcBorders)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


NUMBER_RE = re.compile(r"^[\d\.,%\-\+\(\)\s￥$]+$")


def style_table(table, rules):
    cfg = rules["table"]
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text.strip()
                is_header = ri == 0
                if is_header:
                    align = ALIGN_MAP["center"]
                elif text and NUMBER_RE.match(text):
                    align = ALIGN_MAP[cfg.get("number_alignment", "right")]
                else:
                    align = ALIGN_MAP[cfg.get("text_alignment", "left")]
                p.alignment = align
                for run in p.runs:
                    set_run_font(
                        run,
                        cfg["font_east_asian"],
                        cfg["font_western"],
                        cfg["size_pt"],
                        bold=is_header,
                        color_hex=cfg.get("color_hex"),
                    )
    if cfg.get("three_line_border"):
        apply_three_line_border(table)
    if cfg.get("repeat_header_row") and len(table.rows) > 0:
        set_repeat_header(table.rows[0])


def _add_field_run(paragraph, instr_text):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


def add_page_number_footer(doc, rules):
    cfg = rules.get("footer", {})
    if not cfg.get("enabled"):
        return
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r.text = ""
    p.alignment = ALIGN_MAP.get(cfg.get("alignment", "center"))
    style = cfg.get("style", "page_of_total")
    east = cfg.get("font_east_asian", "宋体")
    size = cfg.get("size_pt", 9)

    if style == "page_of_total":
        p.add_run("第 ")
        _add_field_run(p, " PAGE ")
        p.add_run(" 页 共 ")
        _add_field_run(p, " NUMPAGES ")
        p.add_run(" 页")
    else:
        p.add_run("- ")
        _add_field_run(p, " PAGE ")
        p.add_run(" -")

    for r in p.runs:
        set_run_font(r, east, "Times New Roman", size, color_hex=cfg.get("color_hex"))


def format_document(input_path, output_path, rules):
    doc = Document(input_path)
    set_page_setup(doc, rules)

    paragraphs = doc.paragraphs
    first_nonempty_index = 0
    for i, p in enumerate(paragraphs):
        if p.text.strip():
            first_nonempty_index = i
            break

    report = {"cover": 0, "h1": 0, "h2": 0, "h3": 0, "h4": 0, "body": 0, "skipped_empty": 0, "auto_numbered": 0}

    # Tracks how many auto-numbered headings have been seen at each level
    # since the last heading of a shallower level, so numbering restarts
    # correctly within each parent section (e.g. every top-level section's
    # first un-numbered sub-heading is "（一）", the second is "（二）", ...).
    level_rank = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
    counters = {"h1": 0, "h2": 0, "h3": 0, "h4": 0}

    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text
        if not text.strip():
            report["skipped_empty"] += 1
            continue
        is_list = is_list_paragraph(paragraph)
        if is_list:
            # Bullet/numbered list items are body content by definition, even
            # when their text happens to start with something that looks like
            # a heading number (e.g. a bullet describing "01.碎片化信息文件夹
            # 完全为空..." should not become a fake level-4 heading just
            # because the folder name starts with a digit and a period).
            level, has_own_numbering = None, False
        else:
            style_name = paragraph.style.name if paragraph.style else None
            level, has_own_numbering = classify_heading(text, rules, i, first_nonempty_index, style_name=style_name)
        if level is None:
            apply_body_style(paragraph, rules["body"], is_list=is_list)
            report["body"] += 1
        else:
            if level in level_rank:
                cur_rank = level_rank[level]
                for lv, rank in level_rank.items():
                    if rank > cur_rank:
                        counters[lv] = 0
                if not has_own_numbering:
                    counters[level] += 1
                    prefix = make_heading_number(level, counters[level])
                    if paragraph.runs:
                        paragraph.runs[0].text = prefix + paragraph.runs[0].text
                    else:
                        paragraph.add_run(prefix + text)
                    report["auto_numbered"] += 1
            cfg = _cover_cfg(rules) if level == "cover" else rules["heading_levels"][level]
            apply_heading_style(paragraph, cfg)
            report[level] += 1

    for table in doc.tables:
        style_table(table, rules)

    add_page_number_footer(doc, rules)

    doc.save(output_path)
    return report


def main():
    parser = argparse.ArgumentParser(description="Format a Chinese IPO prospectus-style .docx")
    parser.add_argument("input", help="input .docx path")
    parser.add_argument("output", help="output .docx path")
    parser.add_argument("--rules", default=None, help="path to rules JSON (defaults to assets/default_rules.json next to this script)")
    args = parser.parse_args()

    rules_path = args.rules
    if rules_path is None:
        import os
        rules_path = os.path.join(os.path.dirname(__file__), "..", "assets", "default_rules.json")

    rules = load_rules(rules_path)
    report = format_document(args.input, args.output, rules)

    print("排版完成。段落分类统计：")
    for k, v in report.items():
        print(f"  {k}: {v}")
    print(f"输出文件: {args.output}")


if __name__ == "__main__":
    main()
